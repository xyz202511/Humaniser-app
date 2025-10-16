# app.py
import streamlit as st
from io import BytesIO
from docx import Document
import re
import math
import torch
import numpy as np
import faiss
import torch.nn.functional as F
from collections import Counter
from difflib import SequenceMatcher
from sentence_transformers import SentenceTransformer
from transformers import GPT2LMHeadModel, GPT2TokenizerFast, AutoModelForSeq2SeqLM, AutoTokenizer
from datasketch import MinHash, MinHashLSH

st.set_page_config(page_title="Humaniser", layout="wide")
st.title("Humaniser: Plagiarism & AI Detection & Humanisation Tool")

# ---------------------------
# CONFIGURATION
# ---------------------------
SHINGLE_K = 5
MINHASH_PERM = 128
LSH_THRESHOLD = 0.3
EMBED_MODEL_NAME = "paraphrase-MiniLM-L12-v2"
GPT2_MODEL = "distilgpt2"
HUMANISE_MODEL = "google/flan-t5-base"  # For rewriting
USE_FAISS = True

# ---------------------------
# TEXT UTILITIES
# ---------------------------
def normalize_text(t: str):
    return re.sub(r'\s+', ' ', t).strip()

def simple_tokenize(text: str):
    return re.findall(r'\w+', text.lower())

def get_word_shingles(text: str, k=SHINGLE_K):
    tokens = simple_tokenize(text)
    return {" ".join(tokens[i:i+k]) for i in range(len(tokens) - k + 1)}

def minhash_from_shingles(shingles):
    m = MinHash(num_perm=MINHASH_PERM)
    for s in shingles:
        m.update(s.encode('utf8'))
    return m

def jaccard(set_a, set_b):
    return len(set_a & set_b) / len(set_a | set_b) if set_a or set_b else 0.0

def extract_text_from_docx(file_bytes: BytesIO) -> str:
    try:
        doc = Document(file_bytes)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        st.error(f"Failed to read DOCX: {e}")
        return ""
def chunk_text(text, tokenizer, max_tokens=1024):
    tokens = tokenizer.encode(text)
    chunks = [tokens[i:i+max_tokens] for i in range(0, len(tokens), max_tokens)]
    decoded_chunks = [tokenizer.decode(c, skip_special_tokens=True) for c in chunks]
    return decoded_chunks


# ---------------------------
# PLAGIARISM INDEX
# ---------------------------
class PlagiarismIndex:
    def __init__(self):
        self.doc_texts = {}
        self.doc_shingles = {}
        self.doc_minhashes = {}
        self.lsh = MinHashLSH(threshold=LSH_THRESHOLD, num_perm=MINHASH_PERM)
        self.embed_model = SentenceTransformer(EMBED_MODEL_NAME)
        self.corpus_embeddings = None
        self.corpus_ids = []
        self.faiss_index = None

    def add_document(self, doc_id, text):
        text = normalize_text(text)
        if not text: return
        self.doc_texts[doc_id] = text
        shingles = get_word_shingles(text)
        self.doc_shingles[doc_id] = shingles
        m = minhash_from_shingles(shingles)
        self.doc_minhashes[doc_id] = m
        self.lsh.insert(doc_id, m)

        emb = self.embed_model.encode([text], convert_to_numpy=True)
        if self.corpus_embeddings is None:
            self.corpus_embeddings = emb
        else:
            self.corpus_embeddings = np.vstack([self.corpus_embeddings, emb])
        self.corpus_ids.append(doc_id)
        if USE_FAISS:
            self._build_faiss()

    def _build_faiss(self):
        emb = self.corpus_embeddings.astype('float32')
        faiss.normalize_L2(emb)
        self.faiss_index = faiss.IndexFlatIP(emb.shape[1])
        self.faiss_index.add(emb)

    def query(self, text, top_k=5):
        text = normalize_text(text)
        shingles = get_word_shingles(text)
        q_m = minhash_from_shingles(shingles)
        candidates = self.lsh.query(q_m)

        emb = self.embed_model.encode([text], convert_to_numpy=True).astype('float32')
        faiss.normalize_L2(emb)
        emb_results = []

        if self.faiss_index is not None:
            D, I = self.faiss_index.search(emb, top_k)
            for score, idx in zip(D[0], I[0]):
                if idx < len(self.corpus_ids) and score > 0:
                    emb_results.append(self.corpus_ids[idx])
        else:
            sims = []
            qn = emb[0] / np.linalg.norm(emb[0])
            for i, cid in enumerate(self.corpus_ids):
                v = self.corpus_embeddings[i]
                sims.append((np.dot(qn, v / np.linalg.norm(v)), cid))
            sims = sorted(sims, reverse=True)[:top_k]
            emb_results.extend([cid for _, cid in sims])

        final_ids = list(dict.fromkeys(candidates + emb_results))
        results = []
        for cid in final_ids:
            jac = jaccard(shingles, self.doc_shingles[cid])
            results.append({
                "doc_id": cid,
                "jaccard": jac,
                "text": self.doc_texts[cid]
            })
        return sorted(results, key=lambda x: x["jaccard"], reverse=True)

# ---------------------------
# AI-LIKENESS DETECTOR
# ---------------------------
class AIDetector:
    def __init__(self, model_name=GPT2_MODEL, device=None):
        self.tokenizer = GPT2TokenizerFast.from_pretrained(model_name)
        self.model = GPT2LMHeadModel.from_pretrained(model_name)
        self.model.eval()
        self.device = device or ("cuda" if torch.cuda.is_available() else "cpu")
        self.model.to(self.device)
        for p in self.model.parameters():
            p.requires_grad = False
        self.n_positions = getattr(self.model.config, "n_positions", 1024)

    def _per_token_logprobs(self, enc_ids, chunk_size=512, overlap=64):
        n = enc_ids.size(0)
        if n <= 1: return []
        token_logps = [None]*n
        stride = max(1, chunk_size - overlap)
        for start in range(0, n, stride):
            end = min(start + chunk_size, n)
            chunk = enc_ids[start:end].unsqueeze(0).to(self.device)
            with torch.no_grad():
                logits = self.model(chunk).logits
                log_probs = F.log_softmax(logits, dim=-1)
            L = chunk.size(1)
            if L <= 1: continue
            target_ids = chunk[0, 1:]
            lp = log_probs[0, :-1, :]
            tok_lp_values = lp[torch.arange(L-1), target_ids].cpu().tolist()
            for i, val in enumerate(tok_lp_values):
                global_idx = start + i + 1
                if global_idx < n and token_logps[global_idx] is None:
                    token_logps[global_idx] = val
            if all(x is not None for x in token_logps[1:]): break
        filled = [lp for lp in token_logps[1:] if lp is not None]
        return filled

    def perplexity(self, text, chunk_size=None, overlap=64):
        if not text.strip(): return float("inf")
        enc = self.tokenizer.encode(text, return_tensors="pt")[0]
        if chunk_size is None: chunk_size = min(self.n_positions, 512)
        per_token_logps = self._per_token_logprobs(enc, chunk_size=chunk_size, overlap=overlap)
        if not per_token_logps: return float("inf")
        avg_logp = sum(per_token_logps)/len(per_token_logps)
        ppl = math.exp(-avg_logp)
        return float(ppl)

    def _repetition_score(self, text):
        toks = [t.lower() for t in self.tokenizer.tokenize(text) if t.strip()]
        if not toks: return 0.0
        counts = Counter(toks)
        max_freq = max(counts.values())
        fraction = max_freq/len(toks)
        return min(100.0, fraction*100*2)

    def _ttr_score(self, text):
        toks = [t.lower() for t in self.tokenizer.tokenize(text) if t.strip()]
        if not toks: return 0.0
        ttr = len(set(toks))/len(toks)
        return (1 - ttr)*100

    def ai_score(self, text):
        ppl = self.perplexity(text)
        if math.isinf(ppl): return {"ppl": None, "score": None}
        logppl = math.log(max(ppl,1e-8))
        low, high = math.log(10), math.log(200)
        frac = max(0.0, min(1.0, (logppl-low)/(high-low)))
        perplexity_score = frac*100
        rep = self._repetition_score(text)
        ttr_s = self._ttr_score(text)
        w_ppl, w_rep, w_ttr = 0.6, 0.25, 0.15
        final = perplexity_score*w_ppl + rep*w_rep + ttr_s*w_ttr
        final = max(0.0, min(100.0, final))
        return {"ppl": round(ppl,2), "score": round(final,2),
                "components": {"perplexity": round(perplexity_score,2),
                               "repetition": round(rep,2),
                               "ttr": round(ttr_s,2)}}

# ---------------------------
# Humanisation Model
# ---------------------------
@st.cache_resource
def load_humaniser_model():
    model = AutoModelForSeq2SeqLM.from_pretrained(HUMANISE_MODEL)
    tokenizer = AutoTokenizer.from_pretrained(HUMANISE_MODEL)
    return model, tokenizer

humanise_model, humanise_tokenizer = load_humaniser_model()

def humanise_text(text, max_len=512):
    inputs = humanise_tokenizer.encode("Paraphrase: " + text, return_tensors="pt", truncation=True, max_length=max_len)
    outputs = humanise_model.generate(inputs, max_length=max_len, num_beams=5, early_stopping=True)
    return humanise_tokenizer.decode(outputs[0], skip_special_tokens=True)

# ---------------------------
# MAIN APP
# ---------------------------
uploaded_file = st.file_uploader("Upload DOCX", type=["docx"])
if uploaded_file is not None:
    text = extract_text_from_docx(uploaded_file)
    st.subheader("Original Text")
    st.text_area("Content", text, height=300)

    detector = AIDetector()
    pl_index = PlagiarismIndex()
    pl_index.add_document("uploaded_doc", text)  # Can add more corpus docs later

    if st.button("Check AI-likeness"):
        chunks = chunk_text(text, detector.tokenizer, max_tokens=1024)
        ai_scores = [detector.ai_score(c) for c in chunks]
    
        # Aggregate results
        avg_score = sum([c['score'] for c in ai_scores if c['score'] is not None]) / len(ai_scores)
        st.subheader("AI-likeness Result")
        st.write(f"Average AI-Likeness Score: {avg_score:.2f} / 100")
    
        # Show per-chunk breakdown
        for i, score in enumerate(ai_scores, 1):
            st.write(f"Chunk {i}: {score['score']:.2f} / 100")


    if st.button("Humanise Document"):
        humanised_text = humanise_text(text)
        st.subheader("Humanised Text")
        st.text_area("Humanised Content", humanised_text, height=300)

        # Re-check
        ai_after = detector.ai_score(humanised_text)
        plag_after = pl_index.query(humanised_text)
        st.subheader("After Humanisation")
        st.write(f"AI Score: {ai_after['score']} / 100")
        st.subheader("Plagiarism Matches")
        for r in plag_after[:5]:
            st.markdown(f"- Doc ID: {r['doc_id']}, Similarity: {round(r['jaccard']*100,2)}%")

        # Download
        doc = Document()
        doc.add_paragraph(humanised_text)
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        st.download_button("Download Humanised DOCX", doc_bytes, file_name="Humanised_Output.docx")
